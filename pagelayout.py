import cv2
import numpy as np
import os
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import this

img_path = r'D:\Desktop\PBE\Images\3.jpg'

# WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"

# def set_number_of_columns(section, cols):
#     """ sets number of columns through xpath. """
#     section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(cols))

document = Document()
# document.add_heading("Title", 0)

# para1 = document.add_paragraph("demo add paragraph to document using docx python")
# section = document.sections[0]
# sectPr = section._sectPr
# cols = sectPr.xpath('./w:cols')[0]
# cols.set(qn('w:num'),'2')

# document = Document()
# document.add_paragraph("this page should be normal 1 column")
# document.add_paragraph("this page should have 2 columns and it seems to work ok")

# document = Document('custom_styles.docx')
table = document.add_table(1, 2)
# table.cell(0, 0).text = 'Left Text'
# table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
# table.cell(0, 1).text = 'Right Text'
# table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
row=table.add_row().cells
p=row[0].add_paragraph('left justified text')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=row[1].add_paragraph('right justified text')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER

this.i

document.save('demo.docx')